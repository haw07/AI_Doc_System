import re
from docx import Document

def parse_docx(file_path):
    doc = Document(file_path)
    full_text = []
    
    # 1. Grab standard paragraphs (Header info)
    for p in doc.paragraphs:
        if p.text.strip():
            full_text.append(p.text.strip())
            
    # 2. Grab text from Tables (The meat of your document)
    for table in doc.tables:
        for row in table.rows:
            # Join cell texts with a space to keep "Party A" and "BANK ABC" together
            row_text = " ".join(cell.text.strip() for cell in row.cells)
            full_text.append(row_text)
            
    text = "\n".join(full_text)
    print("Full text extracted:\n", text)  # Debug to see the new layout
    return text

def extract_entities_rule_based(text):
    # Initialize dictionary with the default Business Day value
    entities = {"Calendar": "TARGET"}

    # Patterns updated to use \s+ which matches newlines, and specific anchors
    patterns = {
        "Counterparty": r"Party A\s+(.*)",
        "Initial Valuation Date": r"Initial Valuation Date\s+(.*)",
        "Notional": r"Notional Amount\s*\(N\)\s+(.*)",
        "Valuation Date": r"Valuation Date\s+(.*)",
        "Maturity": r"(?:Termination|Maturity) Date\s+(.*)",
        "Underlying": r"Underlying\s+(.*)",
        "Coupon": r"Coupon\s*\(C\)\s+(.*)",
        "Barrier": r"Barrier\s*\(B\)\s+(.*)" # Matches "Barrier (B)" specifically
    }

    for key, pattern in patterns.items():
        # re.IGNORECASE is used to handle potential casing differences
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            # We split by newline and take the first line to avoid grabbing 
            # the rest of the document if the match is too broad.
            value = match.group(1).split('\n')[0].strip()
            entities[key] = value

    return entities